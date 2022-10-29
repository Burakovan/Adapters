from pathlib import Path
import pathlib
import docx
import markdown
import os
from PyPDF2 import PdfFileReader
import shutil
from constants import PROJECT_ROOT, ASSETS_PATH
from openpyxl import load_workbook
import codecs
import pandas as pd

class Word:
   def __init__(self):
       """
       A path from where a file is taken
       """

       self.path = PROJECT_ROOT / 'all_files'
   def reading(self):
       """
       Opens and reads .docx files
       """

       path_to_data=Path(self.path)
       for file in path_to_data.glob('*.docx'):
           doc = docx.Document(file)
           all_paragraphs = doc.paragraphs
           for paragraph in all_paragraphs:
               print(paragraph.text)

   def writing(self):
       """
       Opens .docx files and adds text in them
       """

       path_to_data = Path(self.path)
       for file in path_to_data.glob('*.docx'):
           doc = docx.Document(file)
           doc.add_paragraph(input())
           doc.save(file)


class Md:
    def __init__(self):
        """
        A path from where a file is taken
        """

        self.path=PROJECT_ROOT / 'all_files'
    def reading(self):
        """
        Opens and reads .md files
        """

        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.md'):
            with open(file,'r',encoding='utf-8') as f:
                text = f.read()
                md_file = markdown.markdown(text)
                print(md_file)
    def writing(self):
        """
        Opens .md files and adds text in them
        """

        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.md'):
            with open(file,'r', encoding='utf-8') as f:
                text = f.read()
                md_file = markdown.markdown(text)
            with open(file, 'w', encoding='utf-8') as f:
                f.write(md_file)
                f.write(input())
                f.close()

class Pdf:
    def __init__(self):
        """
        A path from where a file is taken
        """

        self.path=PROJECT_ROOT / 'all_files'
    def reading(self):
        """
        Opens and reads .pdf files
        """

        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.pdf'):
            file = open(file, 'rb')
            pdfReader = PdfFileReader(file)
            number_of_pages = pdfReader.getNumPages()
            for page_number in range(number_of_pages):
                page = pdfReader.getPage(page_number)
                page_content = page.extractText()
                print(page_content)

    def writing(self):
        """
        Converts .pdf files into .txt files and adds text in them
        """

        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.pdf'):
            reader = PdfFileReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"

            text_file  = open(ASSETS_PATH / 'from_pdf.txt', 'w', encoding='utf-8')
            n = text_file.write(text)
            n = text_file.write(input())
            text_file.close()


class Txt:
    def __init__(self):
        """
        A path from where a file is taken
        """

        self.path = PROJECT_ROOT / 'all_files'

    def reading_txt(self):
        """
        Opens and reads .txt files
        """

        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.txt'):
            file = open(file, encoding='utf-8')
            read = file.readlines()

            modified = ""
            for string in read:
                modified += string.strip()
            print(modified)

    def writing_txt(self):
        """
        Opens .txt files and writes in them
        """

        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.txt'):
            file = open(file, 'a', encoding='utf-8')
            file.write('\n' + input())

class Excel:
    def __init__(self):
        """
        A path from where a file is taken
        """

        self.path = PROJECT_ROOT / 'all_files'

    def reading_excel(self):
        """
        Opens and reads .xlsx files
        """
        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.xlsx'):
            with open('xlsx to txt.txt', 'w') as txt_file:
                pd.read_excel(file).to_string(txt_file, index=False)
            with open('xlsx to txt.txt', 'r') as txt_file:
                for string in txt_file:
                    print(string)

    def writing_excel(self):
        """
        Creates a new sheet in .xlsx file for writing
        """

        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.xlsx'):
            wb = load_workbook(file)

            wb.create_sheet('Новый лист')
            wb.active = wb['Новый лист']
            ws = wb.active
            ws['A1'].value = input('A1: ')
            ws['A2'].value = input('A2: ')
            ws['A3'].value = input('A3: ')
            ws['A4'].value = input('A4: ')

            ws['B1'].value = input('B1: ')
            ws['B2'].value = input('B2: ')
            ws['B3'].value = input('B3: ')
            ws['B4'].value = input('B4: ')

            wb.save(file)


class Html:
    def __init__(self):
        """
        A path from where a file is taken
        """

        self.path = PROJECT_ROOT / 'all_files'

    def reading_html(self):
        """
        Opens and reads .html files
        """

        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.html'):
            f1 = codecs.open(file, 'r', "utf-8")
            print(f1.read())

    def writing_html(self):
        """
        Opens .html files and adds text
        """

        path_to_data = Path(self.path)
        for file in path_to_data.glob('*.html'):
            html_f = open(file, 'w')

            html_upd = '''<!Doctype html>
            <head>
            <title>Warning</title>
            </head>
            <body>
            <h2>Welcome To The Paradise</h2>
    
            <p>Here can live cute Angels.</p>
    
            </body>
            </html>'''

            html_f.write(html_upd)
            html_f.close()

def environment(base_path):
    """
    Creates a folder for changed files
    """

    path = pathlib.Path(base_path)
    path.mkdir(parents=True, exist_ok=True)

def saving():
    """
    Moves changed files to a new folder
    """

    source_path = PROJECT_ROOT / 'all_files'
    allfiles = os.listdir(source_path)
    for file in allfiles:
        src_path = os.path.join(source_path, file)
        dst_path = os.path.join(ASSETS_PATH, file)
        shutil.move(src_path, dst_path)

if __name__=="__main__":
    environment(ASSETS_PATH)
    path_to_files = Path(PROJECT_ROOT / 'all_files')
    if len(os.listdir(path_to_files)) == 0:
        print("Directory is empty")
    elif len(os.listdir(path_to_files)) > 1:
        print ("Only 1 file should be in a directory")
    else:
        for fname in os.listdir('.'):
            if not fname.endswith('.txt') or not fname.endswith('.pdf') \
                    or not fname.endswith('.md') or not fname.endswith('.docx') \
                    or not fname.endswith('.xlsx') or not fname.endswith('.html'):
                print("This type of file is not supported")

        for file in path_to_files.glob('*.md'):
            md_file = Md()
            md_file.reading()
            md_file.writing()
            saving()

        for file in path_to_files.glob('*.docx'):
            doc_file = Word()
            doc_file.reading()
            doc_file.writing()
            saving()

        for file in path_to_files.glob('*.pdf'):
            pdf_file = Pdf()
            pdf_file.reading()
            pdf_file.writing()

        for file in path_to_files.glob('*.txt'):
            txt_file = Txt()
            txt_file.reading_txt()
            txt_file.writing_txt()
            saving()

        for file in path_to_files.glob('*.xlsx'):
            xl_file = Excel()
            xl_file.reading_excel()
            xl_file.writing_excel()
            saving()

        for file in path_to_files.glob('*.html'):
            html_file = Html()
            html_file.reading_html()
            html_file.writing_html()
            saving()


